"""
parser.py — 文件解析器
======================
PDF / DOCX / Markdown / 纯文本 → 统一文本输出。

外部调用：
    result = parse_file("path/to/doc.pdf")
    # → ParseResult(text="...", source_file="doc.pdf", file_type="pdf")

异常处理：
    · 未知扩展名 → UnsupportedFileError
    · 文件不存在 → FileNotFoundError
    · 解析失败   → ParseError（含原始异常）
"""

import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional


# ──────────────────────────────────────────────
# 类型定义
# ──────────────────────────────────────────────

@dataclass
class ParseResult:
    text: str                         # 解析后的纯文本
    source_file: str                  # 原始文件名
    file_type: str                    # 文件扩展名（小写，无点）
    metadata: dict = field(default_factory=dict)  # 可选元信息


class UnsupportedFileError(Exception):
    pass


class ParseError(Exception):
    pass


# ──────────────────────────────────────────────
# 公开入口
# ──────────────────────────────────────────────

def parse_file(file_path: str) -> ParseResult:
    """根据文件扩展名自动路由到对应解析器

    检查顺序：扩展名 → 存在性 → 解析
    这样不支持的格式不需要访问文件系统即可报错。
    """
    path = Path(file_path)
    ext = path.suffix.lower().lstrip(".")

    if ext not in ("pdf", "docx", "md", "markdown", "txt"):
        raise UnsupportedFileError(
            f"Unsupported file type: .{ext} (supported: pdf, docx, md, txt)"
        )

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    DISPATCH = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "md": _parse_markdown,
        "markdown": _parse_markdown,
        "txt": _parse_text,
    }
    return DISPATCH[ext](path)


# ──────────────────────────────────────────────
# PDF 解析（pdfminer.six）
# ──────────────────────────────────────────────

def _parse_pdf(path: Path) -> ParseResult:
    try:
        from pdfminer.high_level import extract_text
    except ImportError:
        raise ImportError(
            "pdfminer.six not installed. Run: pip install pdfminer.six"
        )

    try:
        text = extract_text(str(path))
        if not text.strip():
            raise ParseError("PDF extracted no text (possibly scanned document)")

        return ParseResult(
            text=_normalize(text),
            source_file=path.name,
            file_type="pdf",
            metadata={"pages": _count_pdf_pages(str(path))},
        )
    except ParseError:
        raise
    except Exception as e:
        raise ParseError(f"PDF parse failed: {e}") from e


def _count_pdf_pages(file_path: str) -> Optional[int]:
    """尝试获取 PDF 页数（非关键信息，失败返回 None）"""
    try:
        from pdfminer.pdfparser import PDFParser
        from pdfminer.pdfdocument import PDFDocument

        with open(file_path, "rb") as f:
            parser = PDFParser(f)
            doc = PDFDocument(parser)
            return len(list(doc.get_pages()))
    except Exception:
        return None


# ──────────────────────────────────────────────
# DOCX 解析（python-docx）
# ──────────────────────────────────────────────

def _parse_docx(path: Path) -> ParseResult:
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx not installed. Run: pip install python-docx"
        )

    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        # 也提取表格内容
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            text += "\n\n" + "\n".join(rows)

        return ParseResult(
            text=_normalize(text),
            source_file=path.name,
            file_type="docx",
            metadata={
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
            },
        )
    except Exception as e:
        raise ParseError(f"DOCX parse failed: {e}") from e


# ──────────────────────────────────────────────
# Markdown 解析
# ──────────────────────────────────────────────

def _parse_markdown(path: Path) -> ParseResult:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # 备选编码
        text = path.read_text(encoding="gbk")

    # 保留 Markdown 结构信息（标题标记、列表等）
    # chunker 后续会用这些结构做更智能的切片
    return ParseResult(
        text=_normalize(text),
        source_file=path.name,
        file_type="md",
        metadata={
            "headings": _count_md_headings(text),
        },
    )


def _count_md_headings(text: str) -> int:
    """统计 Markdown 标题数量"""
    return len(re.findall(r"^#{1,6}\s+", text, re.MULTILINE))


# ──────────────────────────────────────────────
# 纯文本解析
# ──────────────────────────────────────────────

def _parse_text(path: Path) -> ParseResult:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="gbk")

    return ParseResult(
        text=_normalize(text),
        source_file=path.name,
        file_type="txt",
    )


# ──────────────────────────────────────────────
# 文本清洗
# ──────────────────────────────────────────────

def _normalize(text: str) -> str:
    """统一文本格式：
    - 将 \r\n 统一为 \n
    - 压缩连续空行（最多保留两个 \n）
    - 去除首尾空白
    """
    # Windows → Unix 换行
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # 压缩连续空行：将 3+ 个 \n 缩减为 2 个
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# ──────────────────────────────────────────────
# 批量解析
# ──────────────────────────────────────────────

def parse_files(file_paths: List[str]) -> List[ParseResult]:
    """批量解析多个文件"""
    results = []
    errors: List[tuple] = []

    for fp in file_paths:
        try:
            results.append(parse_file(fp))
        except Exception as e:
            errors.append((fp, e))

    if errors:
        summary = "; ".join(f"{f}: {e}" for f, e in errors)
        print(f"⚠ {len(errors)}/{len(file_paths)} files failed: {summary}")

    return results
